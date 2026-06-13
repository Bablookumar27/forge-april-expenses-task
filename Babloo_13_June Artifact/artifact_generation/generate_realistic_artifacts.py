import os
import random
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageDraw, ImageFont, ImageFilter
import math

out_dir = r"C:\Users\Bablu\.gemini\antigravity\scratch\artifact_generation\artifacts"
os.makedirs(out_dir, exist_ok=True)

# Helper for generic items
GROCERY_ITEMS = ["Milk 1 Gal", "Eggs 1 Dozen", "Whole Wheat Bread", "Chicken Breast 1lb", "White Rice 2lb", "Pasta 16oz", "Tomato Sauce 24oz", "Cheddar Cheese 8oz", "Apples 3lb", "Bananas 1lb", "Orange Juice 64oz", "Butter 16oz", "Yogurt 32oz", "Cereal 18oz", "Oatmeal 15oz", "Peanut Butter 16oz", "Jam 18oz", "Coffee 12oz", "Tea 16oz", "Sugar 4lb"]
OFFICE_ITEMS = ["Pens 12pk", "Pencils 24pk", "Highlighters 4pk", "Sticky Notes 4pk", "Legal Pads 12pk", "Printer Paper 500sht", "Stapler", "Staples 1000pk", "Paper Clips 1000pk", "Binder Clips 100pk", "Tape 3pk", "Scissors", "Ruler", "Calculator", "Whiteboard Markers 4pk", "Eraser", "Whiteout", "Folders 10pk", "Binders 4pk", "Index Cards 100pk"]

def get_random_items(item_list, n=10):
    return random.sample(item_list, n)

print("1. Generating realistic CSVs...")
def generate_csv(filename, target_row, headers, data, omit_header=False):
    df = pd.DataFrame(data, columns=headers)
    df.to_csv(os.path.join(out_dir, filename), index=False, header=not omit_header)

generate_csv("data_04.csv", 2, ["Transaction_ID", "Timestamp", "Description", "Category", "Qty", "Unit_Price", "Total_Price"], 
    [["TRX-991", "2026-04-09 10:15", "Paper Towels", "Supplies", 2, 15.00, 30.00],
     ["TRX-992", "2026-04-10 14:30", "Bulk Snacks", "Food", 5, 29.00, 145.00], # Target
     ["TRX-993", "2026-04-11 09:00", "Bottled Water", "Food", 3, 5.00, 15.00],
     ["TRX-994", "2026-04-12 11:20", "Trash Bags", "Supplies", 1, 20.00, 20.00],
     ["TRX-995", "2026-04-13 16:45", "Cleaning Wipes", "Supplies", 4, 4.50, 18.00]], omit_header=True)

art_data = [["ORD-10" + str(i), f"2026-04-0{i}", item, 1, 10.00, 10.00] for i, item in enumerate(get_random_items(OFFICE_ITEMS, 3), 1)]
art_data.append(["ORD-104", "2026-04-15", "Markers and Canvas", 2, 42.50, 85.00]) # Target
art_data.extend([["ORD-10" + str(i), f"2026-04-1{i}", item, 1, 15.00, 15.00] for i, item in enumerate(get_random_items(OFFICE_ITEMS, 4), 5)])
generate_csv("data_07.csv", 4, ["Order_ID", "Date", "Item_Description", "Quantity", "Unit_Price", "Total"], art_data)

tags_data = [["TX-20" + str(i), f"2026-04-0{i}", item, 5, 2.00, 10.00] for i, item in enumerate(["Pens", "Pads"], 1)]
tags_data.insert(1, ["TX-202", "2026-04-05", "Name Tags and Lanyards", 50, 0.71, 35.50]) # Target
generate_csv("data_15.csv", 2, ["Trans_ID", "Date", "Description", "Qty", "Price", "Total"], tags_data)

med_data = [["RX-30" + str(i), f"2026-04-0{i}", item, 1, 5.00, 5.00] for i, item in enumerate(["Bandages", "Tape", "Ointment", "Ice Pack"], 1)]
med_data.insert(4, ["RX-305", "2026-04-02", "First Aid Kit Supplies", 1, 45.20, 45.20]) # Target
generate_csv("data_19.csv", 5, ["Receipt_ID", "Date", "Item", "Qty", "Unit", "Total"], med_data)

groc_data = [["GR-40" + str(i), f"2026-04-0{i}", item, 1, 3.00, 3.00] for i, item in enumerate(get_random_items(GROCERY_ITEMS, 15), 1)]
generate_csv("data_24.csv", 0, ["ID", "Date", "Product", "Qty", "Price", "Total"], groc_data)

coffee_data = [["CF-50" + str(i), f"2026-04-0{i}", "Black Coffee", 1, 2.50, 2.50] for i in range(1, 10)]
generate_csv("data_29.csv", 0, ["Txn", "Date", "Item", "Qty", "Price", "Total"], coffee_data)

print("2. Generating realistic XLSXs...")
def generate_xlsx(filename):
    wb = Workbook()
    ws = wb.active
    if filename == "data_10.xlsx":
        ws.title = "Order Form"
        ws['A1'] = "STATIONERY ORDER FORM"
        ws['A1'].font = Font(bold=True, size=16)
        ws.append([])
        headers = ["Item ID", "Description", "Qty", "Unit Price", "Total Price"]
        ws.append(headers)
        for i, item in enumerate(get_random_items(OFFICE_ITEMS, 8)):
            ws.append([f"ITM-{i}", item, 1, 5.00, 5.00])
        ws['D12'] = "Date: 2026-04-20"
        ws['E12'] = "Total: $12.50" # Target
        ws['E12'].font = Font(bold=True)
    elif filename == "data_20.xlsx":
        ws.title = "Sheet1"
        ws['A1'] = "THE FORGE - BUDGET TRACKER"
        ws['A1'].font = Font(bold=True, size=16)
        ws.append([])
        # Missing Date column
        headers = ["Transaction ID", "Category", "Allocated", "Spent"]
        ws.append(headers)
        for cell in ws[3]: cell.font = Font(bold=True)
        for i in range(10):
            ws.append([f"TRX-00{i}", "Supplies", 100, 20])
        ws['D14'] = "Remaining Balance"
        ws['D14'].font = Font(bold=True)
        ws['E14'] = "$210.00" # Target
        ws['E14'].font = Font(bold=True)
    wb.save(os.path.join(out_dir, filename))

generate_xlsx("data_10.xlsx")
generate_xlsx("data_20.xlsx")

print("3. Generating realistic DOCXs...")
def generate_docx(filename, title, date, target_item, target_total, note=None):
    doc = Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Date: {date}").bold = True
    
    doc.add_paragraph("Bill To:\nThe Forge\nc/o Megan Sullivan\n123 Education Way, Detroit, MI")
    doc.add_paragraph("Invoice Details:").bold = True
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Line Total'
    
    items = get_random_items(OFFICE_ITEMS, 3)
    for item in items:
        row = table.add_row().cells
        row[0].text = "1"
        row[1].text = item
        row[2].text = "$10.00"
        row[3].text = "$10.00"
        
    # Target Row
    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = target_item
    row[2].text = target_total
    row[3].text = target_total
    
    p2 = doc.add_paragraph(f"\nSubtotal: {target_total}\nTax: $0.00\n")
    p2.add_run(f"TOTAL DUE: {target_total}").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if note:
        doc.add_paragraph(f"\nNote: {note}").italic = True
        
    doc.save(os.path.join(out_dir, filename))

generate_docx("file_02.docx", "Bookstore Invoice", "2026-04-18", "Leadership Workbooks", "$120.00")
generate_docx("file_06.docx", "Museum Ticket Confirmation", "2026-04-22", "Group Tickets", "$65.00")
generate_docx("file_09.docx", "Business Quote", "2026-04-10", "Custom T-Shirts", "$250.00")
generate_docx("file_12.docx", "Deli Catering Invoice", "2026-04-25", "Sandwich Platters", "$185.00")
generate_docx("file_17.docx", "Expense Claim", "2026-04-28", "Extra Snacks", "$24.00", "For the boys who stayed late")
generate_docx("doc_22.docx", "College Bookstore Invoice", "2026-04-05", "Biology 101 Textbook", "$150.00")
generate_docx("doc_27.docx", "Dry Cleaning Invoice", "2026-04-08", "Women's Suits and Blouses", "$45.00")

print("4. Generating Clean Digital PDFs...")
def generate_pdf_table(filename, title, date, target_item, target_total, dup_page=False):
    doc = SimpleDocTemplate(os.path.join(out_dir, filename), pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    elements.append(Paragraph(f"<b>{title}</b>", styles['Title']))
    elements.append(Paragraph(f"Date: {date}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    data = [['Qty', 'Description', 'Unit Price', 'Total']]
    for item in get_random_items(OFFICE_ITEMS, 4):
        data.append(['1', item, '$5.00', '$5.00'])
    data.append(['1', target_item, target_total, target_total])
    
    t = Table(data, colWidths=[50, 300, 80, 80])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ]))
    elements.append(t)
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(f"<b>GRAND TOTAL: {target_total}</b>", styles['Normal']))
    
    if dup_page:
        elements.extend(elements.copy()) # simple dup
        
    doc.build(elements)

generate_pdf_table("file_05.pdf", "Bus Charter Company Invoice", "2026-04-26", "Charter Bus Services", "$350.00")
generate_pdf_table("file_11.pdf", "Receipt: Reusable Water Bottles", "2026-04-14", "Water Bottles Bulk", "$40.00", dup_page=True)
generate_pdf_table("doc_21.pdf", "Sporting Goods Store Receipt", "2026-04-11", "Youth Basketball Shoes & Gear", "$120.00")
generate_pdf_table("doc_25.pdf", "Auto Repair Shop Invoice", "2026-04-07", "Oil Change and Tire Rotation", "$85.00")
generate_pdf_table("doc_26.pdf", "Event Registration Confirmation", "2026-04-01", "Riverfront Spring Classic 5K", "$35.00")
generate_pdf_table("doc_28.pdf", "Ticketmaster Export", "2026-04-10", "Detroit Lions Football Game Tickets", "$200.00")

print("5. Generating Messy Physical Receipt PDFs...")
def generate_messy_receipt(filename, vendor, date, target_item, target_total, handwritten=None, blur=False, skew=0, orient=0):
    # Create thermal paper background
    img = Image.new('RGB', (600, 800), color='#f4f4f4')
    d = ImageDraw.Draw(img)
    
    # Try to load a monospace font, fallback to default
    try:
        font = ImageFont.truetype("consola.ttf", 20)
        bold_font = ImageFont.truetype("consolab.ttf", 24)
        hw_font = ImageFont.truetype("comic.ttf", 28)
    except:
        font = ImageFont.load_default()
        bold_font = font
        hw_font = font
        
    receipt_text = f"""
    *********************************
           {vendor}
           123 Main Street
           Detroit, MI 48227
    *********************************
    Date: {date}
    Receipt #: 492819
    ---------------------------------
    1  Filler Item 1          $ 5.00
    2  Filler Item 2          $ 8.50
    1  {target_item:<20} {target_total}
    ---------------------------------
    Subtotal:               {target_total}
    Tax (6%):               $ 0.00
    ---------------------------------
    TOTAL:                  {target_total}
    
    Visa ending in 4421
    Auth Code: 0092A
    *********************************
    """
    
    d.text((50, 50), receipt_text, fill='#333333', font=font)
    
    if handwritten:
        d.text((100, 600), handwritten, fill='blue', font=hw_font)
        
    # Apply transformations
    if blur:
        img = img.filter(ImageFilter.GaussianBlur(1.5))
    
    if skew != 0:
        img = img.rotate(skew, fillcolor='#f4f4f4', expand=True)
        
    if orient != 0:
        img = img.rotate(orient, fillcolor='#f4f4f4', expand=True)
        
    pdf_path = os.path.join(out_dir, filename)
    img.save(pdf_path, "PDF", resolution=100.0)

generate_messy_receipt("file_01.pdf", "PIZZA RESTAURANT", "2026-04-12", "Large Pizzas", "$45.20", skew=3)
generate_messy_receipt("file_03.pdf", "CLOTHING STORE", "2026-04-09", "Winter Parka", "$65.00", handwritten="For the new boys", orient=180)
generate_messy_receipt("file_08.pdf", "PRINTING SHOP", "2026-04-16", "Program Flyers", "$28.00", blur=True)
generate_messy_receipt("file_13.pdf", "AV RENTAL CO", "2026-04-24", "Projector and Screen", "$75.00", skew=-4)
generate_messy_receipt("file_14.pdf", "CASH RECEIPT", "2026-04-15", "Guest Speaker Fee", "$100.00", handwritten="Paid in cash", blur=True)
generate_messy_receipt("file_16.pdf", "CAFETERIA", "2026-04-21", "Lunch Account Deposit", "$50.00", handwritten="For the boys")
generate_messy_receipt("file_18.pdf", "TRANSIT KIOSK", "2026-04-18", "Emergency Bus Passes", "$15.00", handwritten="Given to student A", skew=5)
generate_messy_receipt("doc_23.pdf", "MUSIC STORE", "2026-04-06", "Choir Sheet Music", "$22.00", skew=7)
generate_messy_receipt("doc_30.pdf", "BUTCHER SHOP", "2026-04-24", "Bratwurst & Thighs", "$42.00", blur=True)

print("Realistic generation complete!")
